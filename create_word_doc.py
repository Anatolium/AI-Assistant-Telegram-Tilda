from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_technical_doc():
    doc = Document()
    # Заголовок
    title = doc.add_heading('Важные технические особенности реализации бота', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 1. OpenAI API и библиотеки
    doc.add_heading('1. OpenAI API и библиотеки', 1)
    # Версии библиотек
    doc.add_heading('Версии библиотек', 2)
    code = doc.add_paragraph()
    code.add_run('''# requirements.txt
openai==1.12.0  # Важно! Новая версия библиотеки
python-telegram-bot==20.8
flask==3.0.2
flask-cors==4.0.0
google-api-python-client==2.118.0
google-auth-httplib2==0.2.0
google-auth-oauthlib==1.2.0''').font.name = 'Courier New'
    # Правильное подключение OpenAI Assistant
    doc.add_heading('Правильное подключение OpenAI Assistant', 2)
    code = doc.add_paragraph()
    code.add_run('''from openai import OpenAI

client = OpenAI(api_key="your_api_key")

# Создание ассистента (делается один раз)
assistant = client.beta.assistants.create(
    name="Sheets_google_bot",
    instructions="Ваш промпт...",
    model="gpt-4-turbo-preview"  # Используем новейшую модель
)

# Использование существующего ассистента
assistant = client.beta.assistants.retrieve("asst_2Vpr26lbYlgLjML6CH6fbb64")

# Создание треда
thread = client.beta.threads.create()

# Отправка сообщения
message = client.beta.threads.messages.create(
    thread_id=thread.id,
    role="user",
    content=user_message
)

# Запуск ассистента
run = client.beta.threads.runs.create(
    thread_id=thread.id,
    assistant_id=assistant.id
)

# Получение ответа (с ожиданием завершения)
while True:
    run_status = client.beta.threads.runs.retrieve(
        thread_id=thread.id,
        run_id=run.id
    )
    if run_status.status == 'completed':
        break
    time.sleep(1)

# Получение сообщений
messages = client.beta.threads.messages.list(thread_id=thread.id)''').font.name = 'Courier New'
    # 2. Настройка CORS для Flask
    doc.add_heading('2. Настройка CORS для Flask', 1)
    # Базовая настройка CORS
    doc.add_heading('Базовая настройка CORS', 2)
    code = doc.add_paragraph()
    code.add_run('''from flask import Flask
from flask_cors import CORS

app = Flask(__name__)

# Правильная настройка CORS для всех роутов
CORS(app, resources={
    r"/*": {
        "origins": "*",
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

# ИЛИ для конкретного роута
CORS(app, resources={
    r"/website-chat": {
        "origins": ["https://your-website.com"],
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type"]
    }
})''').font.name = 'Courier New'
    # Обработка CORS для веб-виджета
    doc.add_heading('Обработка CORS для веб-виджета', 2)
    code = doc.add_paragraph()
    code.add_run('''@app.route('/website-chat', methods=['POST', 'OPTIONS'])
def website_chat():
    if request.method == 'OPTIONS':
        # Правильная обработка префлайт-запросов
        response = make_response()
        response.headers.add('Access-Control-Allow-Origin', '*')
        response.headers.add('Access-Control-Allow-Headers', 'Content-Type')
        response.headers.add('Access-Control-Allow-Methods', 'POST')
        return response
        
    # Обработка основного POST запроса
    data = request.json
    # ... ваш код ...
    return jsonify({"response": assistant_response})''').font.name = 'Courier New'
    # HTML для веб-виджета
    doc.add_heading('HTML для веб-виджета', 2)
    code = doc.add_paragraph()
    code.add_run('''<!-- working_chat_widget.html -->
<script>
fetch('https://your-ngrok-url/website-chat', {
    method: 'POST',
    headers: {
        'Content-Type': 'application/json',
        // Важно! Не добавляйте лишних заголовков
    },
    body: JSON.stringify({
        message: userMessage
    })
})
.then(response => response.json())
.then(data => {
    // Обработка ответа
})
.catch(error => console.error('Error:', error));
</script>''').font.name = 'Courier New'
    # 3. Особенности работы с ngrok
    doc.add_heading('3. Особенности работы с ngrok', 1)
    # Настройка ngrok.yml
    doc.add_heading('Настройка ngrok.yml', 2)
    code = doc.add_paragraph()
    code.add_run('''authtoken: ваш_токен
tunnels:
  flask:
    addr: 5000
    proto: http
    # Важно! Для продакшена используйте постоянный домен
    # domain: your-domain.ngrok.io''').font.name = 'Courier New'
    # Обработка изменения URL
    doc.add_heading('Обработка изменения URL', 2)
    code = doc.add_paragraph()
    code.add_run('''def update_webhook_url(ngrok_url):
    # Сохраняем URL в файл
    with open('webhook_url.txt', 'w') as f:
        f.write(ngrok_url)
    
    # Обновляем вебхук в Telegram
    telegram_url = f"https://api.telegram.org/bot{TOKEN}/setWebhook"
    webhook_url = f"{ngrok_url}/telegram-webhook"
    response = requests.post(telegram_url, json={'url': webhook_url})
    
    # Проверяем успешность
    if response.status_code == 200:
        print("✅ Webhook успешно обновлен")
    else:
        print("❌ Ошибка обновления webhook")''').font.name = 'Courier New'
    # 4. Важные моменты при деплое
    doc.add_heading('4. Важные моменты при деплое', 1)
    # SSL сертификат
    doc.add_heading('SSL сертификат', 2)
    code = doc.add_paragraph()
    code.add_run('''# Для продакшена используйте SSL
if __name__ == '__main__':
    app.run(ssl_context='adhoc')  # Требуется pip install pyOpenSSL''').font.name = 'Courier New'
    # Обработка ошибок API
    doc.add_heading('Обработка ошибок API', 2)
    code = doc.add_paragraph()
    code.add_run('''try:
    response = client.beta.threads.runs.create(...)
except openai.APIError as e:
    print(f"OpenAI API Error: {e}")
    # Возвращаем пользователю понятное сообщение
    return "Извините, произошла ошибка. Попробуйте позже."''').font.name = 'Courier New'
    # Логирование
    doc.add_heading('Логирование', 2)
    code = doc.add_paragraph()
    code.add_run('''import logging

logging.basicConfig(
    filename='bot.log',
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

logger = logging.getLogger(__name__)''').font.name = 'Courier New'
    # 5. Проверка работоспособности
    doc.add_heading('5. Проверка работоспособности', 1)
    # Тест CORS
    doc.add_heading('Тест CORS', 2)
    code = doc.add_paragraph()
    code.add_run('''curl -X OPTIONS -H "Origin: http://localhost:3000" \
     -H "Access-Control-Request-Method: POST" \
     -H "Access-Control-Request-Headers: Content-Type" \
     http://localhost:5000/website-chat -v''').font.name = 'Courier New'
    # Тест OpenAI
    doc.add_heading('Тест OpenAI', 2)
    code = doc.add_paragraph()
    code.add_run('''# Проверка подключения к OpenAI
try:
    assistant = client.beta.assistants.retrieve(ASSISTANT_ID)
    print("✅ OpenAI Assistant API успешно инициализирован")
except Exception as e:
    print(f"❌ Ошибка инициализации OpenAI: {e}")''').font.name = 'Courier New'
    # Сохраняем документ
    doc.save('Технические_детали_бота.docx')


if __name__ == '__main__':
    create_technical_doc()
